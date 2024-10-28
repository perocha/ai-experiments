from docx import Document
from openai import AzureOpenAI
import json
import time

def load_credentials(filename):
    with open(filename) as cred_file:
        return json.load(cred_file)

# Function to call OpenAI for summarizing the session
def generate_summary(input_text):
    return openai_api_call("Generate a summary", input_text, "The goal is to provide a comprehensive yet concise overview of the session.")

# Function to call OpenAI for extracting key takeaways
def generate_key_takeaways(input_text):
    return openai_api_call("Generate key takeaways", input_text, "Extract key takeaways that are practical for families.")

# Function to call OpenAI for providing practical advice
def generate_practical_advice(input_text):
    return openai_api_call("Provide practical advice", input_text, "Provide practical advice that can help families manage daily challenges.")

# Core function to make OpenAI API calls with error handling and retries
def openai_api_call(prompt_title, input_text, prompt_instruction):
    summary = None
    retry_attempts = 0
    while summary is None:
        try:
            print(f"{prompt_title} (Attempt {retry_attempts + 1})...")
            response = openai.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You are an AI expert in Rett Syndrome, specialized on generating summaries of scientific sessions."},
                    {
                        "role": "user", 
                        "content": f"""
                        {prompt_instruction}
                        The output format should be plain text without any markdown or formatting.

                        {input_text}
                        """
                    }
                ],
                max_tokens=500,
                temperature=0.3
            )
            summary = response.choices[0].message.content
            print(f"{prompt_title} completed successfully.")
        except openai.error.RateLimitError as e:
            retry_after = int(e.headers.get("Retry-After", 8))
            print(f"Rate limit reached. Retrying in {retry_after} seconds...")
            time.sleep(retry_after)
        except openai.error.OpenAIError as e:
            print(f"Error generating {prompt_title}: {e}. Retrying in 5 seconds...")
            time.sleep(5)
        except Exception as e:
            print(f"Unexpected error for {prompt_title}: {e}. Retrying in 5 seconds...")
            time.sleep(5)
        retry_attempts += 1
    return summary




# Load the document
document = Document("mydoc.docx")

# Initialize an empty list for sessions and variables for the current session
sessions = []
current_session = None
current_section = None

# Define the section headers we're looking for
sections = {
    "Author information:": "authorInfo",
    "Affiliations:": "affiliations",
    "Abstract:": "abstract",
    "Session notes:": "sessionNotes",
    "Transcript:": "transcript"
}

# Iterate through paragraphs in the document
for paragraph in document.paragraphs:
    if paragraph.style.style_id == "Heading2":
        if current_session is not None:
            sessions.append(current_session)
        current_session = {
            "title": paragraph.text,
            "authorInfo": "",
            "affiliations": "",
            "abstract": "",
            "sessionNotes": "",
            "transcript": ""
        }
        current_section = None
    elif any(paragraph.text.startswith(header) for header in sections.keys()):
        for header, key in sections.items():
            if paragraph.text.startswith(header):
                current_section = key
                current_session[current_section] = paragraph.text.replace(header, "").strip()
                break
    elif current_section is not None and current_session is not None:
        current_session[current_section] += paragraph.text + "\n"

# Append the last session if it exists
if current_session is not None:
    sessions.append(current_session)

print(f"Found {len(sessions)} sessions in the document.")






# Load OpenAI credentials
openai_keys_file = "../credentials_openai_azure.json"
openai_model_info = load_credentials(openai_keys_file)
openai = AzureOpenAI(
    api_key=openai_model_info["api_key"],
    api_version=openai_model_info["model_version"],
    azure_endpoint=openai_model_info["endpoint"],
    azure_deployment=openai_model_info["deployment_id"]
)

# Create a new Word document for output
output_doc = Document()
output_doc.add_heading("Rett Syndrome World Congress", level=1)

# Loop through each session to generate summaries, key takeaways, and practical advice
for i, session in enumerate(sessions):
    print(f"\nProcessing session {i + 1}/{len(sessions)}: {session['title']}")

    # Prepare the input text for the session
    input_text = (
        f"Author information: {session['authorInfo']}\n"
        f"Affiliations: {session['affiliations']}\n"
        f"Abstract: {session['abstract']}\n"
        f"Transcript: {session['transcript']}\n"
    )

    # Call each prompt function
    summary = generate_summary(input_text)
    key_takeaways = generate_key_takeaways(input_text)
    practical_advice = generate_practical_advice(input_text)

    # Add session details to the document
    output_doc.add_heading(session["title"], level=2)
    output_doc.add_heading("Author Information", level=3)
    output_doc.add_paragraph(session["authorInfo"])
    output_doc.add_heading("Affiliations", level=3)
    output_doc.add_paragraph(session["affiliations"])

    # Add results from each prompt
    output_doc.add_heading("Summary", level=3)
    output_doc.add_paragraph(summary)
    output_doc.add_heading("Key Takeaways", level=3)
    output_doc.add_paragraph(key_takeaways)
    output_doc.add_heading("Practical Advice", level=3)
    output_doc.add_paragraph(practical_advice)

# Save the output document
output_filename = "Rett_Syndrome_World_Congress_Summary.docx"
output_doc.save(output_filename)

print(f"\nDocument saved as {output_filename}")
