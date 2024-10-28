from docx import Document
from openai import AzureOpenAI
import json

def load_credentials(filename):
    with open(filename) as cred_file:
        return json.load(cred_file)


# Load the document
document = Document("mydoc.docx")

# Initialize an empty list for sessions and variables for the current session
sessions = []
current_session = None
current_section = None

# Define the section headers we're looking for
sections = {
    "Author information:": "authorInfo",
    "Affiliations:": "affiliations",
    "Abstract:": "abstract",
    "Session notes:": "sessionNotes",
    "Transcript:": "transcript"
}

# Iterate through paragraphs in the document
for paragraph in document.paragraphs:
    # Check if paragraph is a new session heading
    if paragraph.style.style_id == "Heading2":
        # If there is a current session, append it to the sessions list
        if current_session is not None:
            sessions.append(current_session)
        
        # Start a new session dictionary
        current_session = {
            "title": paragraph.text,
            "authorInfo": "",
            "affiliations": "",
            "abstract": "",
            "sessionNotes": "",
            "transcript": ""
        }
        current_section = None
    
    # Check if the paragraph text starts with any of the defined section headers
    elif any(paragraph.text.startswith(header) for header in sections.keys()):
        # Set the current section based on the detected header
        for header, key in sections.items():
            if paragraph.text.startswith(header):
                current_section = key
                # Remove the header text to only store the content
                current_session[current_section] = paragraph.text.replace(header, "").strip()
                break
    
    # Append content to the current section
    elif current_section is not None and current_session is not None:
        # Append text with a new line for readability
        current_session[current_section] += paragraph.text + "\n"

# Append the last session if it exists
if current_session is not None:
    sessions.append(current_session)

# Display session titles and their structured content for verification
print(f"Found {len(sessions)} sessions")


session_id = 0
# Concatenate the necessary session details to provide full context for summarization
input_str = (
    f"Author information: {sessions[session_id]['authorInfo']}\n"
    f"Affiliations: {sessions[session_id]['affiliations']}\n"
    f"Abstract: {sessions[session_id]['abstract']}\n"
    f"Transcript: {sessions[session_id]['transcript']}\n"
)

openai_keys_file = "../credentials_openai_azure.json"
openai_model_info = load_credentials(openai_keys_file)
# gets the API Key from environment variable AZURE_OPENAI_API_KEY
openai = AzureOpenAI(
    api_key=openai_model_info["api_key"],
    api_version=openai_model_info["model_version"],
    azure_endpoint=openai_model_info["endpoint"],
    azure_deployment=openai_model_info["deployment_id"]
)

try:
    response = openai.chat.completions.create(
        model="gpt-4",
        messages=[
            {"role": "system", "content": "You are an AI expert in Rett Syndrome, specialized on generating summaries of scientific sessions."},
            {
            "role": "user", 
            "content": f"""
            
            Summarize the following transcript of a session presented at the Rett Syndrome World Congress.
            The goal is to provide a comprehensive yet concise overview of the session.
            Please ensure the summary is suitable for individuals with a general knowledge of Rett Syndrome,
            focusing on aspects that would be meaningful for families.
            The output format should be plain text without any markdown or formatting.

            {input_str}
            """
            }
        ],
        max_tokens=500,
        temperature=0.3
    )

except openai.error.RateLimitError as e:
    retry_after = int(e.headers.get("Retry-After", 8))
    print(f"Rate limited. Retrying in {retry_after} seconds.")
    #time.sleep(retry_after)
    exit()
except openai.error.OpenAIError as e:
    print(f"Error generating analysis: {e}")
except Exception as e:
    print(f"Unexpected error: {e}")


# Extract the summary from the response
summary = response.choices[0].message.content

# Create a new Word document for output
output_doc = Document()

# Add a title for the document
output_doc.add_heading("Rett Syndrome World Congress", level=1)

# Add the author information, affiliations, and abstract to the document
output_doc.add_heading("Author Information", level=2)
output_doc.add_paragraph(sessions[session_id]["authorInfo"])
output_doc.add_heading("Affiliations", level=2)
output_doc.add_paragraph(sessions[session_id]["affiliations"])

# Add session title
session_title = sessions[session_id]["title"]
output_doc.add_heading(session_title, level=2)

# Add summary section
output_doc.add_heading("Summary", level=3)
output_doc.add_paragraph(summary)

# Save the output document
output_filename = "Rett_Syndrome_World_Congress_Summary.docx"
output_doc.save(output_filename)

print(f"Summary and key takeaways have been saved to {output_filename}")
