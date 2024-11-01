import sys
from docx import Document
from openai import AzureOpenAI
import json
import time
import argparse

def load_credentials(filename):
    with open(filename) as cred_file:
        return json.load(cred_file)



# Function to load a prompt from a text file
def load_prompt(filename):
    with open(filename, 'r') as file:
        return file.read().strip()

# Load each prompt from its respective file
curate_prompt = load_prompt("prompt_curate_transcript.txt")
summary_prompt = load_prompt("prompt_summary.txt")
key_takeaways_prompt = load_prompt("prompt_key_takeaways.txt")
practical_advice_prompt = load_prompt("prompt_practical_advice.txt")



# Function to call OpenAI for summarizing the session
def generate_curated_transcript(input_text, language, max_tokens):
    return openai_api_call("Generate a curated transcript", input_text, curate_prompt, language, max_tokens)

def generate_summary(input_text, language, max_tokens):
    return openai_api_call("Generate a summary", input_text, summary_prompt, language, max_tokens)

# Function to call OpenAI for extracting key takeaways
def generate_key_takeaways(input_text, language, max_tokens):
    return openai_api_call("Generate key takeaways", input_text, key_takeaways_prompt, language, max_tokens)

# Function to call OpenAI for providing practical advice
def generate_practical_advice(input_text, language, max_tokens):
    return openai_api_call("Provide practical advice", input_text, practical_advice_prompt, language, max_tokens)



# Core function to make OpenAI API calls with error handling and retries
def openai_api_call(prompt_title, input_text, prompt_instruction, language, max_tokens):
    summary = None
    retry_attempts = 0
    while summary is None:
        try:
            print(f"{prompt_title} (Attempt {retry_attempts + 1})...")
            response = openai.chat.completions.create(
                model="gpt-4-32k",
                messages=[
                    {"role": "system", "content": "You are an AI expert in Rett Syndrome."},
                    {
                        "role": "user", 
                        "content": f"""
                        {prompt_instruction}
                        The output format should be plain text without any markdown or formatting.
                        The output should be in {language}.

                        The input text is as follows:
                        {input_text}
                        """
                    }
                ],
                max_tokens=max_tokens,
                temperature=0.2
            )
            summary = response.choices[0].message.content
            print(f"{prompt_title} completed successfully.")
        except openai.error.RateLimitError as e:
            retry_after = int(e.headers.get("Retry-After", 8))
            print(f"Rate limit reached. Retrying in {retry_after} seconds...")
            time.sleep(retry_after)
        except openai.error.OpenAIError as e:
            print(f"Error generating {prompt_title}: {e}. Retrying in 5 seconds...")
            time.sleep(5)
        except Exception as e:
            print(f"Unexpected error for {prompt_title}: {e}. Retrying in 5 seconds...")
            time.sleep(5)
        retry_attempts += 1
    return summary


# Load the document
document = Document("mydoc.docx")

# Initialize an empty list for sessions and variables for the current session
sessions = []
current_session = None
current_section = None

# Define the section headers we're looking for
sections = {
    "Author information:": "authorInfo",
    "Affiliations:": "affiliations",
    "Abstract:": "abstract",
    "Session notes:": "sessionNotes",
    "Transcript:": "transcript"
}

# Iterate through paragraphs in the document
for paragraph in document.paragraphs:
    if paragraph.style.style_id == "Heading2":
        if current_session is not None:
            sessions.append(current_session)
        current_session = {
            "title": paragraph.text,
            "authorInfo": "",
            "affiliations": "",
            "abstract": "",
            "sessionNotes": "",
            "transcript": ""
        }
        current_section = None
    elif any(paragraph.text.startswith(header) for header in sections.keys()):
        for header, key in sections.items():
            if paragraph.text.startswith(header):
                current_section = key
                current_session[current_section] = paragraph.text.replace(header, "").strip()
                break
    elif current_section is not None and current_session is not None:
        current_session[current_section] += paragraph.text + "\n"

# Append the last session if it exists
if current_session is not None:
    sessions.append(current_session)

print(f"Found {len(sessions)} sessions in the document.")




# Load OpenAI credentials
openai_keys_file = "../credentials_openai_azure2.json"
openai_model_info = load_credentials(openai_keys_file)
openai = AzureOpenAI(
    api_key=openai_model_info["api_key"],
    api_version=openai_model_info["model_version"],
    azure_endpoint=openai_model_info["endpoint"],
    azure_deployment=openai_model_info["deployment_id"]
)




# Define the argument parser
parser = argparse.ArgumentParser(description='Process session number, language, and max tokens.')
parser.add_argument('--session_number', type=int, nargs='?', default=1, help='Session number (default: 1)')
parser.add_argument('--language', type=str, default='English', help='Language (default: English)')
parser.add_argument('--max_tokens', type=int, default=100, help='Maximum number of tokens (default: 100)')

# Parse the arguments
args = parser.parse_args()

# Get the session number, language, and max tokens from the arguments
session_number = args.session_number
language = args.language
max_tokens = args.max_tokens

# Check if session number is within the valid range
if session_number < 1 or session_number > len(sessions):
    print(f"Invalid session number: {session_number}. Please choose a number between 1 and {len(sessions)}.")
    sys.exit(1)

# Use the language and max_tokens arguments as needed
print(f"Selected language: {language}")
print(f"Max tokens: {max_tokens}")







# Create a new Word document for output
output_doc = Document()
output_doc.add_heading("Rett Syndrome World Congress", level=1)

# Process only the specified session
session = sessions[session_number - 1]  # Convert to 0-based index
print(f"\nProcessing session {session_number}/{len(sessions)}: {session['title']}")

# Prepare the input text for the session
input_text = (
    f"Author information: {session['authorInfo']}\n"
    f"Affiliations: {session['affiliations']}\n"
    f"Abstract: {session['abstract']}\n"
    f"Transcript: {session['transcript']}\n"
)

# Call each prompt function
max_tokens = 5000
curated_transcript = generate_curated_transcript(input_text, language, max_tokens)
print("Curated transcript generated successfully")
print(curated_transcript)

# Save the curated transcript and summary to a text file
with open("curated_transcript.txt", "w") as file:
    file.write(curated_transcript)

print("\n************************************\n")
max_tokens = 5000
summary = generate_summary(curated_transcript, language, max_tokens)
with open("summary.txt", "w") as file:
    file.write(summary)

max_tokens = 5000
key_takeaways = generate_key_takeaways(curated_transcript, language, max_tokens)
#practical_advice = generate_practical_advice(curated_transcript)

# Add session details to the document
output_doc.add_heading(session["title"], level=2)
#output_doc.add_heading("Author Information", level=3)
#output_doc.add_paragraph(session["authorInfo"])
#output_doc.add_heading("Affiliations", level=3)
#output_doc.add_paragraph(session["affiliations"])

# Add results from each prompt
output_doc.add_heading("Summary", level=3)
output_doc.add_paragraph(summary)
output_doc.add_heading("Key Takeaways", level=3)
output_doc.add_paragraph(key_takeaways)
#output_doc.add_heading("Practical Advice", level=3)
#output_doc.add_paragraph(practical_advice)

# Save the output document
output_filename = f"Rett_Syndrome_World_Congress_Session_{session_number}_Summary.docx"
output_doc.save(output_filename)

print(f"\nDocument saved as {output_filename}")
